"""
Genera Documentacion-Endpoints-Portal-Paciente.docx con la doc de los
10 endpoints implementados en F1..F8 (F9 y F10 son cambios de modelo).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COLOR_PRIMARIO = RGBColor(0x1F, 0x3A, 0x68)
COLOR_METODO = {
    "GET":    RGBColor(0x1B, 0x6E, 0x4A),
    "POST":   RGBColor(0xB8, 0x6E, 0x00),
    "PUT":    RGBColor(0x1F, 0x57, 0xA8),
    "DELETE": RGBColor(0xB0, 0x2A, 0x2A),
}


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = COLOR_PRIMARIO
    return h


def add_method_line(doc, metodo, ruta):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{metodo}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = COLOR_METODO.get(metodo, RGBColor(0, 0, 0))
    r2 = p.add_run(ruta)
    r2.font.name = "Consolas"
    r2.font.size = Pt(11)
    r2.bold = True
    return p


def add_kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(10)
    # Ancho de columnas
    for row in t.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.5)
    return t


def add_code_block(doc, code, lang_label=None):
    if lang_label:
        p = doc.add_paragraph()
        r = p.add_run(lang_label)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    # Sombreado gris claro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    return p


def add_errors(doc, items):
    p = doc.add_paragraph()
    r = p.add_run("Errores")
    r.bold = True
    for code, desc in items:
        bp = doc.add_paragraph(style="List Bullet")
        rc = bp.add_run(f"{code}")
        rc.bold = True
        rc.font.color.rgb = COLOR_METODO["DELETE"]
        bp.add_run(f"  {desc}")


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A68")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Documento ─────────────────────────────────────────────────
doc = Document()

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Portada
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Documentación de Endpoints\nPortal del Paciente")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = COLOR_PRIMARIO

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Policlínico Parroquial San José — Backend (Express + MongoDB)\nIteración 2 · Tareas F1–F10 · 2026-04-30")
rs.font.size = Pt(12)
rs.italic = True

doc.add_paragraph()

# ── Información general ──
add_heading(doc, "Información general", 1)
add_kv_table(doc, [
    ("Base URL",       "http://localhost:3000"),
    ("Prefijo común",  "/api/paciente"),
    ("Autenticación",  "Bearer JWT (Authorization: Bearer <token>) + rol PACIENTE"),
    ("Middlewares",    "verifyToken → requirePaciente (validan token, rol y pacienteId vinculado)"),
    ("CORS",           "Origen permitido: http://localhost:5173"),
    ("Content-Type",   "application/json (excepto avatar: multipart/form-data)"),
    ("Errores comunes", "401 token inválido · 403 rol incorrecto · 500 error servidor"),
])

doc.add_paragraph()
add_heading(doc, "Resumen de endpoints", 2)
tabla = doc.add_table(rows=1, cols=3)
tabla.style = "Light Grid Accent 1"
hdr = tabla.rows[0].cells
for i, h in enumerate(["Método", "Ruta", "Descripción"]):
    hdr[i].text = ""
    rh = hdr[i].paragraphs[0].add_run(h)
    rh.bold = True
endpoints_resumen = [
    ("GET",    "/api/paciente/me",                                "Obtiene el perfil del paciente autenticado"),
    ("PUT",    "/api/paciente/me",                                "Actualiza datos personales (DNI inmutable)"),
    ("PUT",    "/api/paciente/me/password",                       "Cambia la contraseña del portal"),
    ("POST",   "/api/paciente/me/avatar",                         "Sube/actualiza foto de perfil"),
    ("GET",    "/api/paciente/citas",                             "Historial de citas con filtros y paginación"),
    ("GET",    "/api/paciente/ordenes",                           "Órdenes de laboratorio/imagen del paciente"),
    ("GET",    "/api/paciente/terminos",                          "Términos y condiciones vigentes"),
    ("GET",    "/api/paciente/notificaciones",                    "Lista paginada de notificaciones"),
    ("PUT",    "/api/paciente/notificaciones/marcar-todas-leidas","Marca todas las notificaciones como leídas"),
    ("PUT",    "/api/paciente/notificaciones/:id/leer",           "Marca una notificación específica como leída"),
    ("DELETE", "/api/paciente/notificaciones/:id",                "Elimina (soft) una notificación"),
]
for m, r, d in endpoints_resumen:
    row = tabla.add_row().cells
    row[0].text = ""
    rm = row[0].paragraphs[0].add_run(m)
    rm.bold = True
    rm.font.color.rgb = COLOR_METODO[m]
    row[1].text = ""
    rr = row[1].paragraphs[0].add_run(r)
    rr.font.name = "Consolas"
    rr.font.size = Pt(9.5)
    row[2].text = d

doc.add_page_break()


# ── F1: GET /me ──
add_heading(doc, "F1 · GET /api/paciente/me", 1)
add_method_line(doc, "GET", "/api/paciente/me")
doc.add_paragraph("Devuelve los datos personales del paciente autenticado.")
add_kv_table(doc, [
    ("Auth",   "Bearer JWT + rol PACIENTE"),
    ("Body",   "—"),
    ("Query",  "—"),
])
add_code_block(doc, """{
  "success": true,
  "data": {
    "_id": "65f1...",
    "nombres": "Carlos",
    "apellidos": "Vargas",
    "dni": "12345678",
    "telefono": "+51987654321",
    "correo": "carlos@example.com",
    "fechaNacimiento": "1990-05-12T00:00:00.000Z",
    "sexo": "M",
    "estadoCivil": "SOLTERO",
    "direccion": "Av. Lima 123",
    "distrito": "San Isidro",
    "apoderadoNombre": "",
    "apoderadoParentesco": "",
    "apoderadoTelefono": "",
    "avatar": "/uploads/avatares/65f1....jpg",
    "edad": 35
  }
}""", "Respuesta 200 OK")
add_errors(doc, [("404", "Paciente no encontrado"), ("500", "Error del servidor")])
divider(doc)


# ── F2: PUT /me ──
add_heading(doc, "F2 · PUT /api/paciente/me", 1)
add_method_line(doc, "PUT", "/api/paciente/me")
doc.add_paragraph("Actualiza los datos personales del paciente. El DNI es inmutable. Si cambia el correo, se sincroniza con la cuenta Usuario para que el paciente siga entrando al portal con el correo nuevo.")
add_kv_table(doc, [
    ("Auth",            "Bearer JWT + rol PACIENTE"),
    ("Content-Type",    "application/json"),
    ("Validaciones",    "Nombres/Apellidos máx 100. Teléfono Perú (+51XXXXXXXXX o 9XXXXXXXX). Correo formato válido. Sexo M/F. Dirección máx 200, Distrito máx 100. Correo único entre Pacientes y Usuarios."),
])
add_code_block(doc, """{
  "nombres": "Carlos",
  "apellidos": "Vargas",
  "telefono": "+51987654321",
  "correo": "carlos@example.com",
  "fechaNacimiento": "1990-05-12",
  "sexo": "M",
  "direccion": "Av. Lima 123",
  "distrito": "San Isidro",
  "apoderadoNombre": "",
  "apoderadoParentesco": "",
  "apoderadoTelefono": ""
}""", "Body")
add_code_block(doc, """{
  "success": true,
  "message": "Perfil actualizado correctamente",
  "data": { ...mismo formato que GET /me... }
}""", "Respuesta 200 OK")
add_errors(doc, [
    ("400", "Datos inválidos (formato, longitud, teléfono, correo, fecha, sexo)"),
    ("404", "Paciente no encontrado"),
    ("409", "Correo o teléfono duplicado en otro paciente o usuario"),
    ("500", "Error del servidor"),
])
divider(doc)


# ── F2b: PUT /me/password ──
add_heading(doc, "F2b · PUT /api/paciente/me/password", 1)
add_method_line(doc, "PUT", "/api/paciente/me/password")
doc.add_paragraph("Cambio de contraseña del portal del paciente. Requiere la contraseña actual para confirmar identidad.")
add_kv_table(doc, [
    ("Auth",         "Bearer JWT + rol PACIENTE"),
    ("Content-Type", "application/json"),
    ("Reglas",       "Mínimo 8 caracteres, al menos 1 mayúscula, 1 minúscula, 1 dígito y 1 carácter especial (!@#$%^&*). passwordNueva ≠ passwordActual."),
])
add_code_block(doc, """{
  "passwordActual": "Carlos123!",
  "passwordNueva": "Carlos456@",
  "passwordConfirm": "Carlos456@"
}""", "Body")
add_code_block(doc, """{
  "success": true,
  "message": "Contraseña actualizada correctamente"
}""", "Respuesta 200 OK")
add_errors(doc, [
    ("400", "Campos faltantes, no coinciden, igual a la actual, o no cumple reglas de complejidad"),
    ("400", "La contraseña actual es incorrecta"),
    ("404", "Usuario no encontrado"),
    ("500", "Error del servidor"),
])
divider(doc)


# ── F4: POST /me/avatar ──
add_heading(doc, "F4 · POST /api/paciente/me/avatar", 1)
add_method_line(doc, "POST", "/api/paciente/me/avatar")
doc.add_paragraph("Sube o actualiza la foto de perfil del paciente. La imagen se guarda en el filesystem y queda disponible en /uploads/avatares/{pacienteId}.{ext}.")
add_kv_table(doc, [
    ("Auth",            "Bearer JWT + rol PACIENTE"),
    ("Content-Type",    "multipart/form-data"),
    ("Campo",           "avatar (archivo único)"),
    ("Tipos permitidos","JPG, JPEG, PNG, WebP (validación por magic bytes, no solo MIME)"),
    ("Tamaño máximo",   "2 MB (2 097 152 bytes)"),
    ("Almacenamiento",  "uploads/avatares/{pacienteId}.{ext} — se sobrescribe en cada subida"),
])
add_code_block(doc, """curl -X POST http://localhost:3000/api/paciente/me/avatar \\
  -H "Authorization: Bearer <token>" \\
  -F "avatar=@/ruta/a/foto.jpg" """, "Ejemplo de petición")
add_code_block(doc, """{
  "success": true,
  "message": "Avatar actualizado correctamente",
  "data": {
    "avatarUrl": "/uploads/avatares/507f1f77bcf86cd799439011.jpg"
  }
}""", "Respuesta 200 OK")
add_errors(doc, [
    ("400", "No se recibió archivo / tipo no permitido / tamaño excedido"),
    ("500", "Error al guardar el archivo"),
])
divider(doc)


# ── F5/getMyCitas ──
add_heading(doc, "F5 · GET /api/paciente/citas", 1)
add_method_line(doc, "GET", "/api/paciente/citas")
doc.add_paragraph("Historial de citas del paciente, con filtros y paginación. Excluye datos de pago.")
add_kv_table(doc, [
    ("Auth",   "Bearer JWT + rol PACIENTE"),
    ("Query",  "estado=PENDIENTE|ATENDIDA|CANCELADA|...  desde=YYYY-MM-DD  hasta=YYYY-MM-DD  pagina=N (default 1, porPagina fijo en 20)"),
    ("Orden",  "fecha desc, hora desc"),
])
add_code_block(doc, """{
  "success": true,
  "data": [ { "_id":"...", "fecha":"...", "estado":"...", "doctorId":{ ... }, ... } ],
  "paginacion": { "total": 42, "pagina": 1, "porPagina": 20, "totalPaginas": 3 }
}""", "Respuesta 200 OK")
add_errors(doc, [("500", "Error del servidor")])
divider(doc)


# ── getMyOrdenes ──
add_heading(doc, "F5b · GET /api/paciente/ordenes", 1)
add_method_line(doc, "GET", "/api/paciente/ordenes")
doc.add_paragraph("Órdenes de laboratorio/imagen del paciente con filtros y paginación.")
add_kv_table(doc, [
    ("Auth",   "Bearer JWT + rol PACIENTE"),
    ("Query",  "estado=PENDIENTE|EN_ANALISIS|FINALIZADO|...  tipoOrden=LABORATORIO|IMAGEN  pagina=N"),
    ("Orden",  "fecha desc"),
])
add_code_block(doc, """{
  "success": true,
  "data": [ { "_id":"...", "fecha":"...", "estado":"...", "items":[ ... ] } ],
  "paginacion": { "total": 8, "pagina": 1, "porPagina": 20, "totalPaginas": 1 }
}""", "Respuesta 200 OK")
add_errors(doc, [("500", "Error del servidor")])
divider(doc)


# ── F3: GET /terminos ──
add_heading(doc, "F3 · GET /api/paciente/terminos", 1)
add_method_line(doc, "GET", "/api/paciente/terminos")
doc.add_paragraph("Devuelve los términos y condiciones vigentes del portal del paciente. El contenido está versionado.")
add_kv_table(doc, [
    ("Auth",   "Bearer JWT + rol PACIENTE"),
    ("Body",   "—"),
    ("Query",  "—"),
])
add_code_block(doc, """{
  "success": true,
  "data": {
    "contenido": "TÉRMINOS Y CONDICIONES DEL PORTAL DEL PACIENTE\\nPoliclínico Parroquial San José...",
    "version": "1.0",
    "fechaActualizacion": "2026-04-30"
  }
}""", "Respuesta 200 OK")
add_errors(doc, [("500", "Error del servidor")])
divider(doc)


# ── F5: GET /notificaciones ──
add_heading(doc, "F5c · GET /api/paciente/notificaciones", 1)
add_method_line(doc, "GET", "/api/paciente/notificaciones")
doc.add_paragraph("Lista paginada de notificaciones del paciente. Excluye las eliminadas (soft delete).")
add_kv_table(doc, [
    ("Auth",       "Bearer JWT + rol PACIENTE"),
    ("Query",      "filtro=TODAS|LEIDAS|NO_LEIDAS (default TODAS)  pagina=N (default 1)  porPagina=N (default 20, máx 50)"),
    ("Orden",      "fechaCreacion desc"),
])
add_code_block(doc, """{
  "success": true,
  "data": [
    {
      "id": "507f1f77bcf86cd799439011",
      "titulo": "Tu próxima cita",
      "mensaje": "Tu cita con el Dr. García está programada para el 2026-05-15 a las 10:30",
      "tipo": "CITA",
      "leida": false,
      "fechaCreacion": "2026-04-30T10:00:00.000Z",
      "fechaLectura": null,
      "link": "/paciente/citas/507f1f77bcf86cd799439012",
      "eliminada": false
    }
  ],
  "paginacion": { "total": 45, "pagina": 1, "porPagina": 20, "totalPaginas": 3 }
}""", "Respuesta 200 OK")
add_errors(doc, [("500", "Error del servidor")])
divider(doc)


# ── F7: marcar todas leídas ──
add_heading(doc, "F7 · PUT /api/paciente/notificaciones/marcar-todas-leidas", 1)
add_method_line(doc, "PUT", "/api/paciente/notificaciones/marcar-todas-leidas")
doc.add_paragraph("Marca como leídas todas las notificaciones no leídas y no eliminadas del paciente. Devuelve cuántas fueron actualizadas.")
add_kv_table(doc, [
    ("Auth",   "Bearer JWT + rol PACIENTE"),
    ("Body",   "—"),
    ("Query",  "—"),
    ("Nota",   "Esta ruta debe registrarse antes de :id/leer para evitar colisiones de matching en Express."),
])
add_code_block(doc, """{
  "success": true,
  "message": "Todas las notificaciones han sido marcadas como leídas",
  "actualizadas": 15
}""", "Respuesta 200 OK")
add_errors(doc, [("500", "Error del servidor")])
divider(doc)


# ── F6: marcar leída individual ──
add_heading(doc, "F6 · PUT /api/paciente/notificaciones/:id/leer", 1)
add_method_line(doc, "PUT", "/api/paciente/notificaciones/:id/leer")
doc.add_paragraph("Marca como leída una notificación específica del paciente. Idempotente: si ya estaba leída, devuelve 200 sin cambios.")
add_kv_table(doc, [
    ("Auth",        "Bearer JWT + rol PACIENTE"),
    ("Path param",  "id (ObjectId de la notificación)"),
    ("Cambios BD",  "leida=true, fechaLectura=ahora"),
])
add_code_block(doc, """{
  "success": true,
  "message": "Notificación marcada como leída"
}""", "Respuesta 200 OK")
add_errors(doc, [
    ("404", "Notificación no encontrada (id inválido, no existe o está eliminada)"),
    ("403", "La notificación pertenece a otro paciente"),
    ("500", "Error del servidor"),
])
divider(doc)


# ── F8: DELETE notificación ──
add_heading(doc, "F8 · DELETE /api/paciente/notificaciones/:id", 1)
add_method_line(doc, "DELETE", "/api/paciente/notificaciones/:id")
doc.add_paragraph("Elimina (soft delete) una notificación del paciente. No la borra físicamente — solo marca eliminada=true.")
add_kv_table(doc, [
    ("Auth",        "Bearer JWT + rol PACIENTE"),
    ("Path param",  "id (ObjectId de la notificación)"),
    ("Cambios BD",  "eliminada=true (mantiene historial)"),
])
add_code_block(doc, """{
  "success": true,
  "message": "Notificación eliminada"
}""", "Respuesta 200 OK")
add_errors(doc, [
    ("404", "Notificación no encontrada (id inválido, no existe o ya está eliminada)"),
    ("403", "La notificación pertenece a otro paciente"),
    ("500", "Error del servidor"),
])
divider(doc)


# ── F9 / F10: Modelos ──
doc.add_page_break()
add_heading(doc, "F9 · Modelo Notificacion", 1)
doc.add_paragraph("Archivo: src/models/Notificacion.ts")
add_kv_table(doc, [
    ("Campos",    "pacienteId (ObjectId, ref Paciente), titulo (max 100), mensaje (max 500), tipo enum [CITA, EXAMEN, RECETA, SISTEMA], leida (default false), fechaCreacion (default now), fechaLectura (nullable), link, eliminada (default false)"),
    ("Índices",   "Compuesto (pacienteId, leida) para búsquedas rápidas + (fechaCreacion desc) para ordenamiento"),
    ("Métodos estáticos", "crearNotificacion(pacienteId, titulo, mensaje, tipo, link?) · obtenerPorPaciente(pacienteId, filtro, pagina, porPagina)"),
    ("Métodos de instancia", "marcarComoLeida() — actualiza leida y fechaLectura solo si no estaba leída"),
])

add_heading(doc, "F10 · Campo avatar en Paciente", 1)
doc.add_paragraph("Archivo: src/models/Paciente.ts")
add_kv_table(doc, [
    ("Campo",     "avatar?: string (default null)"),
    ("Formato",   "Ruta relativa /uploads/avatares/{pacienteId}.{ext}"),
    ("Servido",   "Express static en app.ts: app.use('/uploads', express.static(...))"),
    ("Validación","Magic bytes en controller (FF D8 FF para JPG, 89 50 4E 47 para PNG, RIFF + WEBP) — no solo MIME"),
])


# ── Notas finales ──
doc.add_paragraph()
add_heading(doc, "Notas de seguridad y diseño", 1)
notas = [
    "Todas las rutas pasan por verifyToken + requirePaciente. El JWT incluye pacienteId vinculado.",
    "El DNI no es modificable desde el portal — solo administración lo cambia.",
    "Cuando el paciente cambia su correo, se sincroniza con Usuario para que el login del portal siga funcionando.",
    "Las notificaciones usan soft delete: nunca se eliminan físicamente, solo se marcan eliminada=true.",
    "El avatar valida magic bytes (no solo extensión ni Content-Type) para evitar subir archivos maliciosos disfrazados.",
    "Nombre del archivo de avatar = pacienteId, así una nueva subida sobrescribe el archivo previo y no se acumulan huérfanos.",
    "Si el paciente cambia de extensión (sube un PNG después de un JPG), el archivo anterior con la extensión antigua se elimina del disco.",
    "Las rutas de notificaciones se registran en orden: marcar-todas-leidas antes que :id/leer para evitar que :id capture el literal.",
]
for n in notas:
    doc.add_paragraph(n, style="List Bullet")


import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Documentacion-Endpoints-Portal-Paciente.docx")
doc.save(out)
print("OK ->", out)
